"""
Word DOCX generator structuring native paragraph arrays securely building NEMA templates implicitly.
Expanded to professional 20-100 page depth.
"""

from io import BytesIO
import boto3
from django.conf import settings
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACRONYMS = {
    "EMCA": "Environmental Management and Coordination Act",
    "ESIA": "Environmental and Social Impact Assessment",
    "ESMMP": "Environmental and Social Management and Monitoring Plan",
    "NEMA": "National Environment Management Authority",
    "OSHA": "Occupational Safety and Health Act",
    "WB": "World Bank",
    "WHO": "World Health Organization",
    "NMT": "Non-Motorized Transport",
    "WRMA": "Water Resources Management Authority",
    "PPE": "Personal Protective Equipment",
    "CHA": "Critical Habitat Assessment",
    "BMP": "Biodiversity Management Plan",
    "DMU": "Discrete Management Unit",
    "KBA": "Key Biodiversity Area",
    "IBAT": "Integrated Biodiversity Assessment Tool",
    "IUCN": "International Union for Conservation of Nature",
    "RAP": "Resettlement Action Plan",
    "TMP": "Traffic Management Plan",
    "CBD": "Central Business District",
    "KII": "Key Informant Interviews",
    "FGD": "Focused Group Discussions",
    "SEA": "Strategic Environmental Assessment"
}

def generate_docx_report(project_id: str, tenant_id: str, version: int, report_data: dict) -> tuple:
    """
    Iterates explicit arrays natively mapping dictionary outputs onto nested python-docx constructs natively.
    Expanded for professional depth.
    """
    doc = Document()
    
    # Global Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # 1. COVER PAGE
    doc.add_heading('Environmental Impact Assessment', 0)
    p = doc.add_paragraph(report_data['project']['name'])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(28)
    p.runs[0].bold = True
    
    doc.add_paragraph(f"Project Type: {report_data['project']['type']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Risk Category: {report_data['project']['category']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Estimated Cost: {report_data['project']['investment_value']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"NEMA License Fee: {report_data['project']['nema_fee']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Scale: {report_data['project']['scale_ha']} ha").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Location: {report_data['project']['location_coords']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Technical Fidelity Alert (NEW)
    if not report_data['project'].get('region_data_captured', True):
        warning = doc.add_paragraph()
        run = warning.add_run("\n⚠️ TECHNICAL DATA FIDELITY WARNING: No specific regional baseline data captured for these coordinates. This report is grounded in a NATIONAL PROXY. Lead Expert MUST conduct immediate field verification for soil, hydrology, and biodiversity parameters.")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n" * 4)
    doc.add_paragraph(f"Prepared by: {report_data['project']['lead_consultant']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Consultant Reg: {report_data['project']['consultant_reg']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {report_data['project']['date']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 2. TABLE OF CONTENTS (Placeholder text)
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph("1. Executive Summary\n2. Project Description\n3. Methodology\n4. Baseline Environment\n5. Legislative Framework\n6. Impacts & Mitigations\n7. Public Participation\n8. ESMP Matrix\n9. Conclusion")
    doc.add_page_break()

    # 2. ACRONYMS
    doc.add_heading('Acronyms', level=1)
    table = doc.add_table(rows=0, cols=2)
    for acr, desc in sorted(ACRONYMS.items()):
        row = table.add_row().cells
        row[0].text = acr
        row[1].text = desc
    doc.add_page_break()

    # 3. EXEC SUMMARY
    doc.add_heading('1. Executive Summary', level=1)
    if report_data.get('executive_summary'):
        doc.add_paragraph(report_data['executive_summary'])
    else:
        doc.add_paragraph(f"This report analyzes the {report_data['project']['name']} project. Baseline Grade: {report_data['baseline']['sensitivity_grade']}.")
    
    # Compliance Alerts
    if report_data['baseline'].get('compliance_alerts'):
        doc.add_heading('Compliance Alerts', level=2)
        for alert in report_data['baseline']['compliance_alerts']:
            p = doc.add_paragraph()
            run = p.add_run(f"[{alert['level']}] {alert['message']}")
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
            doc.add_paragraph(f"Remedy: {alert['remedy']}")
    
    # 4. CRITICAL HABITAT ASSESSMENT (NEW)
    if report_data['baseline'].get('cha_flag', {}).get('is_critical'):
        doc.add_heading('Critical Habitat Assessment (CHA)', level=2)
        cha = report_data['baseline']['cha_flag']
        p = doc.add_paragraph()
        run = p.add_run("CRITICAL HABITAT TRIGGERED")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        
        doc.add_paragraph(f"Discrete Management Unit (DMU): {cha['dm_unit']}")
        doc.add_paragraph("Reasons for Designation:").bold = True
        for reason in cha['reasons']:
            doc.add_paragraph(reason, style='List Bullet')
            
        doc.add_heading('Biodiversity Management Plan (BMP) Summary', level=2)
        doc.add_paragraph("Based on IFC PS6 requirements, a net-gain outcome is mandatory for this project.")
        doc.add_paragraph(f"Ecosystem Services: {report_data['baseline'].get('ecosystem_services', 'N/A')}")

    # 4.1 NON-TECHNICAL SUMMARY (SWAHILI)
    if report_data.get('swahili_summary'):
        doc.add_heading('Muhtasari wa Mradi (Non-Technical Summary)', level=2)
        doc.add_paragraph(report_data['swahili_summary'])
        
    # 5.3 LEGAL FRAMEWORK
    doc.add_heading('Legislative and Regulatory Framework', level=1)
    if report_data.get('legal_narrative'):
        doc.add_paragraph(report_data['legal_narrative'])
    else:
        doc.add_paragraph("The project is governed by EMCA 1999 and the Environmental (Impact Assessment and Audit) Regulations, 2003.")

    # Mandatory Permits Table
    if report_data.get('compliance_package'):
        doc.add_heading('Mandatory Statutory Permits', level=2)
        doc.add_paragraph("Based on the sectoral and regional scope of this project, the following permits must be secured prior to site mobilization:")
        
        cp = report_data['compliance_package']
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        row = table.add_row().cells
        row[0].text = "Permit Description"
        row[1].text = "Issuing Authority"
        
        for p in cp.get('permits', []):
            row = table.add_row().cells
            row[0].text = p
            row[1].text = "NEMA/County/Sector Lead Agency"
        
        doc.add_paragraph("\nNote: Failure to comply with these statutory requirements may result in project stoppage per EMCA Section 118.")
        
    # 5. METHODOLOGY
    doc.add_heading('3. Study Methodology', level=1)
    if report_data.get('methodology'):
        doc.add_paragraph(report_data['methodology'])
    else:
        doc.add_paragraph("The study followed EMCA 1999 scoping guidelines and multi-criteria decision analysis.")
        
    doc.add_heading('2. Project Description', level=1)
    if report_data.get('project_description'):
        doc.add_paragraph(report_data['project_description'])
    else:
        doc.add_paragraph(f"The project is a {report_data['project']['type']} development covering {report_data['project']['scale_ha']} hectares.")
    
    doc.add_heading('2.1 Analysis of Alternatives', level=2)
    if report_data.get('alternatives'):
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Alternative'
        hdr[1].text = 'Environmental Impact'
        hdr[2].text = 'Rationale'
        for alt in report_data['alternatives']:
            row = table.add_row().cells
            row[0].text = alt['alternative']
            row[1].text = alt['env_impact']
            row[2].text = alt['rationale']
    else:
        doc.add_paragraph("In accordance with NEMA regulations, the following alternatives were considered:")
        doc.add_paragraph("The 'No Project' Alternative:", style='List Bullet').bold = True
        doc.add_paragraph("This option would maintain the status quo. While avoiding environmental disturbance, it would deny the local community the socioeconomic benefits of infrastructure development and improved service delivery.")
        doc.add_paragraph("The 'Relocation' Alternative:", style='List Bullet').bold = True
        doc.add_paragraph("Shifting the project to an alternative site was evaluated. However, the current site is optimized for proximity to existing utility corridors and community needs.")
        doc.add_paragraph("Alternative Materials and Design:", style='List Bullet').bold = True
        doc.add_paragraph("Utilization of sustainable materials and labor-intensive construction methods is prioritized to maximize local employment and minimize carbon footprint.")

    # 5.2 HYDROGEOLOGICAL ASSESSMENT (For Boreholes)
    if report_data['baseline'].get('hydrogeology'):
        doc.add_heading('Hydrogeological Assessment', level=1)
        doc.add_paragraph("This section details the subsurface conditions and water resource potential within the Lake Victoria South Basin context, adhering to WRA permitting standards.")
        
        h_data = report_data['baseline']['hydrogeology']
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        row = table.add_row().cells
        row[0].text = "Aquifer Type"
        row[1].text = h_data['aquifer_type']
        
        row = table.add_row().cells
        row[0].text = "Target Drilling Depth"
        row[1].text = f"{h_data['target_depth_m']} Meters"
        
        row = table.add_row().cells
        row[0].text = "Expected Safe Yield"
        row[1].text = f"{h_data['expected_yield_m3h']} m³/hour"
        
        row = table.add_row().cells
        row[0].text = "Recharge Potential"
        row[1].text = h_data['recharge_potential']
        
        row = table.add_row().cells
        row[0].text = "Pump Test Protocl"
        row[1].text = h_data['pump_test_duration']
        
        doc.add_paragraph("\nGroundwater extraction will be managed to ensure sustainability and minimize drawdown interference with existing community wells within a 1km radius.")

    # 5. BASELINE ENVIRONMENT
    doc.add_heading('3. Description of Environment', level=1)
    
    if report_data['baseline'].get('historical_context'):
        doc.add_heading('3.1 Historical Regional Context', level=2)
        doc.add_paragraph(report_data['baseline']['historical_context'])
        doc.add_heading('3.2 Climate Data', level=2)
    else:
        doc.add_heading('3.1 Climate Data', level=2)
        
    # Climate Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Month'
    hdr_cells[1].text = 'Temp Avg'
    hdr_cells[2].text = 'Precip (mm)'
    hdr_cells[3].text = 'Humidity (%)'
    
    for m in report_data['baseline']['climate'].get('monthly', []):
        row = table.add_row().cells
        row[0].text = m['month']
        row[1].text = f"{m['temperature']['mean_avg']}°C"
        row[2].text = str(m['precipitation_mm'])
        row[3].text = f"{m['humidity_percent']['max_avg']}%"

    doc.add_page_break()

    # 6. IMPACTS
    doc.add_heading('4. Impacts and Mitigations', level=1)
    for pred in report_data['predictions']:
        doc.add_heading(f"Impact: {pred['category']}", level=2)
        doc.add_paragraph(f"Severity: {pred['severity']} | Probability: {pred['probability']*100:.1f}%")
        doc.add_paragraph(f"Description: {pred['description']}")
        doc.add_paragraph("Mitigations:", style='List Bullet')
        for m in pred['mitigations']:
            doc.add_paragraph(m, style='List Bullet 2')

    doc.add_page_break()

    # 7. PUBLIC PARTICIPATION
    doc.add_heading('5. Public Participation', level=1)
    doc.add_paragraph("Public consultation is a mandatory requirement under EMCA 1999 (revised 2015). The following methods were employed to capture community feedback:")
    doc.add_paragraph("Public Barazas:", style='List Bullet').bold = True
    doc.add_paragraph("Formal community meetings chaired by local administration (Chiefs/Assistant Chiefs) to disclose project details and capture collective concerns.")
    doc.add_paragraph("Questionnaires & Interviews:", style='List Bullet').bold = True
    doc.add_paragraph("Structured questionnaires distributed to project-affected persons (PAPs) and businesses within a 5km radius.")
    doc.add_paragraph("Key Informant Interviews (KII):", style='List Bullet').bold = True
    doc.add_paragraph("Direct engagement with government officials, utility providers (Kenya Power, CWSB), and community leaders.")

    doc.add_heading('5.1 Consultation Findings', level=2)
    doc.add_paragraph(f"Total Submissions Recorded: {report_data['community']['total_count']}")
    
    # Consultation Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Date'
    hdr[1].text = 'Stakeholder / Location'
    hdr[2].text = 'Sentiment'
    hdr[3].text = 'Primary Concerns / Feedback'
    
    for f in report_data['community']['entries']:
        row = table.add_row().cells
        row[0].text = f['date']
        row[1].text = f"{f.get('submitter_name', 'Stakeholder')} - {f['location']}"
        row[2].text = f['sentiment']
        row[3].text = f['text']

    doc.add_page_break()

    # 8. ESMP
    doc.add_heading('6. Environmental and Social Management Plan (ESMP)', level=1)
    doc.add_paragraph("The following matrix outlines the mitigation measures, monitoring indicators, and responsibilities for each project phase.")
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Impact'
    hdr_cells[2].text = 'Mitigation Measure'
    hdr_cells[3].text = 'Monitoring Indicator'
    hdr_cells[4].text = 'Responsibility (Expert / Local Authority)'
    hdr_cells[5].text = 'Cost'

    for item in report_data.get('esmp_table', []):
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.get('phase', ''))
        row_cells[1].text = str(item.get('impact', ''))
        row_cells[2].text = str(item.get('measure', ''))
        row_cells[3].text = str(item.get('indicator', ''))
        row_cells[4].text = str(item.get('resp', ''))
        row_cells[5].text = str(item.get('cost', ''))

    # 8.1 AUXILIARY PLANS
    doc.add_heading('7. Hazard Management and Decommissioning', level=1)
    doc.add_heading('7.1 Hazard Management Plan', level=2)
    doc.add_paragraph(report_data.get('hazard_plan', 'Standard safety protocols apply.'))
    
    doc.add_heading('7.2 Decommissioning Phase', level=2)
    doc.add_paragraph(report_data.get('decommissioning', 'Site restoration following EMCA guidelines.'))

    doc.add_page_break()

    # 9. COMPLIANCE APPENDIX
    doc.add_heading('Appendix A: Compliance Audit', level=1)
    doc.add_paragraph(f"Compliance Score: {report_data['audit']['score']}% | Grade: {report_data['audit']['grade']}")
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Regulation'
    hdr[1].text = 'Status'
    hdr[2].text = 'Evidence/Remedy'
    
    for item in report_data['audit']['passed'] + report_data['audit']['failed']:
        row = table.add_row().cells
        row[0].text = item['regulation_id']
        row[1].text = 'PASSED' if item.get('status') == 'passed' else 'FAILED'
        row[2].text = f"{item['evidence']} {item.get('remedy', '')}"

    doc.add_page_break()

    # 10. DECLARATION PAGES (NEMA MANDATORY)
    doc.add_heading('Certificate of Declaration', level=1)
    
    doc.add_heading('Declaration by the Expert', level=2)
    doc.add_paragraph(
        f"I, {report_data['project']['lead_consultant']}, a registered EIA/Audit Expert (Reg No: {report_data['project']['consultant_reg']}), "
        "hereby certify that the particulars given in this report are correct and true to the best of my knowledge and belief. "
        "I have carried out the assessment in accordance with the Environmental Management and Coordination Act (EMCA) 1999 "
        "and the Environmental (Impact Assessment and Audit) Regulations 2003."
    )
    
    # Professional Certification Block (Signature & Stamp)
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Left Cell: Signature
    left_cell = sig_table.rows[0].cells[0]
    if report_data['project'].get('consultant_signature'):
        try:
            run = left_cell.paragraphs[0].add_run()
            run.add_picture(report_data['project']['consultant_signature'], width=Inches(1.8))
        except Exception as e:
            left_cell.add_paragraph("Signature: ___________________________")
    else:
        left_cell.add_paragraph("Signature: ___________________________")
    
    # Right Cell: Official Stamp
    right_cell = sig_table.rows[0].cells[1]
    if report_data['project'].get('consultant_stamp'):
        try:
            run = right_cell.paragraphs[0].add_run()
            run.add_picture(report_data['project']['consultant_stamp'], width=Inches(2.0))
        except Exception as e:
            pass
    
    doc.add_paragraph(f"Practicing License No: {report_data['project']['consultant_reg']}")
    doc.add_paragraph(f"Date: {report_data['project']['date']}")

    doc.add_paragraph("\n" * 2)
    
    doc.add_heading('Declaration by the Proponent', level=2)
    doc.add_paragraph(
        f"I, ___________________________ (on behalf of {report_data['project']['proponent']['name']}), "
        "submit this Environmental Impact Assessment Report for the proposed project. "
        "I confirm that the project will be implemented in accordance with the mitigation measures "
        "and management plans outlined in this report."
    )
    
    doc.add_paragraph("\n" * 2)
    doc.add_paragraph("Signature: ___________________________")
    doc.add_paragraph(f"Designation: ___________________________")
    doc.add_paragraph(f"PIN No: {report_data['project']['proponent']['pin']}")
    doc.add_paragraph(f"Date: ___________________________")

    # Save
    file_stream = BytesIO()
    doc.save(file_stream)
    docx_bytes = file_stream.getvalue()
    file_size = len(docx_bytes)

    s3_key = f"reports/{tenant_id}/{project_id}/v{version}/report.docx"
    bucket_name = getattr(settings, "AWS_STORAGE_BUCKET_NAME", "ecosense-dummy-bucket")
    storage_success = False

    if getattr(settings, "AWS_ACCESS_KEY_ID", None):
        try:
            s3 = boto3.client('s3')
            s3.put_object(Bucket=bucket_name, Key=s3_key, Body=docx_bytes, ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            url = s3.generate_presigned_url('get_object', Params={'Bucket': bucket_name, 'Key': s3_key}, ExpiresIn=604800)
            storage_success = True
        except Exception as s3_err:
            print(f"S3 Upload failed ({s3_err}). Falling back to local storage.")
    
    if not storage_success:
        # Save to local media for development download support
        from pathlib import Path
        local_path = Path(settings.MEDIA_ROOT) / s3_key
        local_path.parent.mkdir(parents=True, exist_ok=True)
        with open(local_path, 'wb') as f:
            f.write(docx_bytes)
            
        # Returning the relative media URL for the internal DownloadView
        url = f"/media/{s3_key}"
        
    return s3_key, file_size, url
